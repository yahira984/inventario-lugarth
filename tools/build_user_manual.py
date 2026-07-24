from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageChops
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
ASSET_DIR = OUTPUT_DIR / ".manual_assets"
OUTPUT_PATH = OUTPUT_DIR / "Manual_de_Usuario_Inventario_Lugarth.docx"
LOGO_PATH = ROOT / "public" / "images" / "logo.png"
LOGO_CROPPED_PATH = ASSET_DIR / "logo_manual.png"

BLUE = "0B63CE"
DARK_BLUE = "063B72"
NAVY = "082B4C"
LIGHT_BLUE = "EAF4FF"
PALE_BLUE = "F6FAFF"
GREEN = "0E9F6E"
LIGHT_GREEN = "E8F8F1"
RED = "D93025"
LIGHT_RED = "FDEDEC"
ORANGE = "C76B00"
LIGHT_ORANGE = "FFF3E0"
PURPLE = "6D3BD1"
LIGHT_PURPLE = "F2ECFF"
TEAL = "087F8C"
LIGHT_TEAL = "E6F7F8"
GRAY = "5F6B7A"
LIGHT_GRAY = "F3F6F9"
BORDER = "CAD7E5"
WHITE = "FFFFFF"
BLACK = "142033"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_run_font(run, size=None, bold=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=8, color=GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualiza esta tabla en Word: clic derecho > Actualizar campo."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_char_sep, placeholder, fld_char_end])


def prepare_logo() -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.open(LOGO_PATH).convert("RGB")
    background = Image.new("RGB", image.size, image.getpixel((0, 0)))
    diff = ImageChops.difference(image, background)
    bbox = diff.getbbox()
    if bbox:
        left, top, right, bottom = bbox
        image = image.crop((left, top, right, bottom))
    image.save(LOGO_CROPPED_PATH, optimize=True, quality=88)
    return LOGO_CROPPED_PATH


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 12),
        ("Subtitle", 14, GRAY, 0, 8),
        ("Heading 1", 17, DARK_BLUE, 18, 8),
        ("Heading 2", 13.5, BLUE, 13, 6),
        ("Heading 3", 11.5, NAVY, 10, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(10.5)
    styles["List Number"].font.name = "Calibri"
    styles["List Number"].font.size = Pt(10.5)


def configure_sections(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("INVENTARIO LUGARTH  |  MANUAL DE USUARIO")
    set_run_font(run, size=8, bold=True, color=GRAY)
    p.paragraph_format.space_after = Pt(0)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_cover(document: Document, logo_path: Path) -> None:
    band = document.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(band, 6.86)
    cell = band.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=80, bottom=80, start=80, end=80)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CONTROL DE INVENTARIO, ALMACÉN Y EQUIPOS")
    set_run_font(run, size=9, bold=True, color=WHITE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after = Pt(12)
    inline_shape = p.add_run().add_picture(str(logo_path), width=Inches(4.7))
    inline_shape._inline.docPr.set("descr", "Logotipo de Lugarth")
    inline_shape._inline.docPr.set("title", "Lugarth")

    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Manual de usuario")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Inventario Lugarth")
    set_run_font(run, size=19, bold=True, color=BLUE)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(
        "Operación de inventario, entradas, salidas, equipos, compras, "
        "catálogos, reportes y administración."
    )
    set_run_font(run, size=11, color=GRAY)

    document.add_paragraph()
    info = document.add_table(rows=3, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(info, 4.75)
    set_table_layout_fixed(info)
    rows = [
        ("Versión del manual", "1.0"),
        ("Fecha de emisión", date.today().strftime("%d/%m/%Y")),
        ("Dirigido a", "Administradores, almacenistas y consultores"),
    ]
    for row_index, (label, value) in enumerate(rows):
        for col_index, text in enumerate((label, value)):
            cell = info.cell(row_index, col_index)
            set_cell_margins(cell, top=75, bottom=75)
            set_cell_shading(cell, LIGHT_BLUE if col_index == 0 else WHITE)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 6, "color": BORDER},
                bottom={"val": "single", "sz": 6, "color": BORDER},
                left={"val": "single", "sz": 6, "color": BORDER},
                right={"val": "single", "sz": 6, "color": BORDER},
            )
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(
                run,
                size=9.5,
                bold=col_index == 0,
                color=DARK_BLUE if col_index == 0 else BLACK,
            )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run("Documento interno de operación")
    set_run_font(run, size=9, bold=True, color=GRAY)
    document.add_page_break()


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_body(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first, rest = text.split(":", 1)
        run = p.add_run(first + ":")
        set_run_font(run, bold=True, color=NAVY)
        run = p.add_run(rest)
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        run = p.add_run(item)
        set_run_font(run)


def add_steps(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.29)
        p.paragraph_format.first_line_indent = Inches(-0.19)
        run = p.add_run(item)
        set_run_font(run)


def add_callout(document: Document, title: str, text: str, kind: str = "info") -> None:
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (LIGHT_GREEN, GREEN),
        "warning": (LIGHT_ORANGE, ORANGE),
        "danger": (LIGHT_RED, RED),
        "purple": (LIGHT_PURPLE, PURPLE),
        "teal": (LIGHT_TEAL, TEAL),
    }
    fill, accent = colors[kind]
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.75)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    set_cell_border(
        cell,
        left={"val": "single", "sz": 18, "color": accent},
        top={"val": "single", "sz": 4, "color": fill},
        bottom={"val": "single", "sz": 4, "color": fill},
        right={"val": "single", "sz": 4, "color": fill},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=10.2, bold=True, color=accent)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=BLACK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, sum(widths))
    set_table_layout_fixed(table)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = header_row.cells[index]
        cell.width = Inches(width)
        set_cell_shading(cell, DARK_BLUE)
        set_cell_margins(cell, top=80, bottom=80, start=85, end=85)
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=8.5, bold=True, color=WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for col_index, (text, width) in enumerate(zip(row_data, widths)):
            cell = row.cells[col_index]
            cell.width = Inches(width)
            set_cell_shading(cell, WHITE if row_index % 2 == 0 else PALE_BLUE)
            set_cell_margins(cell, top=68, bottom=68, start=85, end=85)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": 4, "color": BORDER},
            )
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            set_run_font(run, size=8.6, color=BLACK)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_screenshot(
    document: Document,
    number: int,
    title: str,
    instruction: str,
    height: float = 1.45,
) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, 6.75)
    row = table.rows[0]
    row.height = Inches(height)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
    set_cell_border(
        cell,
        top={"val": "dashed", "sz": 10, "color": BLUE},
        bottom={"val": "dashed", "sz": 10, "color": BLUE},
        left={"val": "dashed", "sz": 10, "color": BLUE},
        right={"val": "dashed", "sz": 10, "color": BLUE},
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"CAPTURA {number:02d}  |  {title}")
    set_run_font(run, size=10.5, bold=True, color=BLUE)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(instruction)
    set_run_font(run, size=8.8, color=GRAY)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Inserta la imagen en este espacio y elimina este recuadro.")
    set_run_font(run, size=8.3, bold=True, color=DARK_BLUE)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def start_section(document: Document, title: str, intro: str | None = None) -> None:
    document.add_page_break()
    add_heading(document, title, 1)
    if intro:
        add_body(document, intro)


def build_document() -> Document:
    logo_path = prepare_logo()
    document = Document()
    configure_styles(document)
    configure_sections(document)
    core_properties = document.core_properties
    core_properties.title = "Manual de usuario - Inventario Lugarth"
    core_properties.subject = "Operación del sistema de inventario, almacén y equipos"
    core_properties.author = "Lugarth"
    core_properties.keywords = "inventario, almacén, manual, Lugarth"
    core_properties.comments = "Documento interno de operación"

    add_cover(document, logo_path)

    add_heading(document, "Control del documento", 1)
    add_table(
        document,
        ["Dato", "Información"],
        [
            ["Nombre", "Manual de usuario de Inventario Lugarth"],
            ["Versión", "1.0"],
            ["Estado", "Versión operativa para validación y capacitación"],
            ["Responsable", "Administrador del sistema"],
            ["Usuarios", "Administrador, almacenista y consultor"],
            ["Actualización", "Debe revisarse cuando cambie un flujo, permiso o pantalla."],
        ],
        [1.65, 5.1],
    )
    add_callout(
        document,
        "Cómo completar este manual",
        "Los recuadros CAPTURA 01, CAPTURA 02, etc. indican exactamente qué pantalla debe fotografiarse. "
        "Usa datos de prueba, oculta correos o información sensible y conserva el zoom del navegador en 100 %. "
        "Después inserta cada imagen dentro de su recuadro.",
        "info",
    )
    add_heading(document, "Contenido", 1)
    add_table(
        document,
        ["Sección", "Contenido", "Sección", "Contenido"],
        [
            ["1", "Objetivo y alcance", "12", "Catálogos y etiquetas"],
            ["2", "Roles y permisos", "13", "Herramientas y reportes"],
            ["3", "Acceso, registro y sesión", "14", "Administración"],
            ["4", "Navegación y elementos comunes", "15", "Perfil, seguridad y dispositivos"],
            ["5", "Dashboard gerencial", "16", "Alertas y notificaciones"],
            ["6", "Inventario de almacén", "17", "Reglas de stock y trazabilidad"],
            ["7", "Registrar entradas", "18", "Rutinas recomendadas"],
            ["8", "Registrar salidas", "19", "Solución de problemas"],
            ["9", "Devoluciones y mermas", "20", "Glosario"],
            ["10", "Equipos y paquetes", "Anexo A", "Lista maestra de capturas"],
            ["11", "Compras y facturas", "", ""],
        ],
        [0.62, 2.65, 0.72, 2.76],
    )

    start_section(
        document,
        "1. Objetivo y alcance",
        "Este manual explica el uso diario del sistema Inventario Lugarth. Incluye el control de existencias, "
        "entradas y salidas, devoluciones y mermas, equipos o paquetes, facturas XML, proveedores, órdenes de "
        "compra, etiquetas, reportes, usuarios, auditoría y respaldos.",
    )
    add_heading(document, "1.1 Qué controla el sistema", 2)
    add_bullets(
        document,
        [
            "Piezas reales del almacén, con descripción, apodo, categoría, ubicación, marca, costo, proveedor, fotografía y códigos.",
            "Entradas de stock inmediatas para administradores y entradas sujetas a aprobación para almacenistas.",
            "Salidas manuales, por código de barras, QR, cámara o pistola USB.",
            "Devoluciones que regresan stock y mermas que justifican bajas por daño.",
            "Equipos o paquetes formados por varias piezas reales y retiros que descuentan todos sus componentes.",
            "Información de CFDI XML, proveedores, costos y datos SAT.",
            "Trazabilidad de movimientos, acciones de usuarios y respaldos de la base de datos.",
        ],
    )
    add_heading(document, "1.2 Principios de operación", 2)
    add_callout(
        document,
        "Regla principal",
        "El inventario contiene piezas reales. Los equipos o paquetes son recetas que indican cuántas piezas "
        "consume cada equipo. No deben registrarse equipos como si fueran piezas de inventario.",
        "success",
    )
    add_bullets(
        document,
        [
            "Toda entrada, salida, devolución, merma o retiro de equipo debe quedar registrada con información suficiente.",
            "Un almacenista no puede aprobar su propia entrada; el stock se suma hasta que un administrador la revise.",
            "Una orden de compra no cambia el stock. El stock cambia cuando la mercancía entra y la entrada queda aprobada.",
            "No elimines registros para corregir una operación. Usa el módulo correcto y conserva la trazabilidad.",
            "Las sugerencias del identificador visual son apoyo; siempre confirma descripción, medida, número de parte y categoría.",
        ],
    )

    start_section(document, "2. Roles y permisos")
    add_body(
        document,
        "El sistema muestra únicamente las opciones permitidas para cada rol. El primer usuario registrado se convierte "
        "automáticamente en administrador y queda aprobado. Los usuarios posteriores quedan pendientes hasta que un "
        "administrador apruebe su correo y asigne su rol.",
    )
    add_table(
        document,
        ["Módulo o acción", "Administrador", "Almacenista", "Consultor"],
        [
            ["Dashboard gerencial", "Sí", "No", "No"],
            ["Consultar inventario", "Sí", "Sí", "Sí"],
            ["Registrar entradas", "Sí, inmediatas", "Sí, con aprobación", "No"],
            ["Registrar salidas", "Sí", "Sí", "No"],
            ["Devoluciones y mermas", "Sí", "Sí", "No"],
            ["Equipos y retiros", "Sí", "Sí", "No"],
            ["Importar XML", "Sí", "No", "No"],
            ["Aprobar entradas", "Sí", "No", "No"],
            ["Proveedores y órdenes", "Sí", "No", "No"],
            ["Editar/eliminar materiales", "Sí", "No", "No"],
            ["Categorías y códigos", "Sí", "No", "No"],
            ["Identificador visual", "Sí", "Sí", "Sí"],
            ["Centro de reportes", "Sí", "No", "Sí"],
            ["Usuarios, auditoría y respaldos", "Sí", "No", "No"],
            ["Perfil propio", "Sí", "Sí", "Sí"],
        ],
        [2.75, 1.35, 1.35, 1.3],
    )
    add_callout(
        document,
        "Protección del administrador",
        "Un administrador no debe quitarse a sí mismo la aprobación ni cambiar su propio rol. Tampoco debe eliminarse "
        "la última cuenta administradora, porque el sistema quedaría sin una persona capaz de aprobar usuarios.",
        "danger",
    )

    start_section(document, "3. Acceso, registro y sesión")
    add_heading(document, "3.1 Iniciar sesión", 2)
    add_steps(
        document,
        [
            "Abre la dirección del sistema proporcionada por la empresa.",
            "Escribe el correo registrado.",
            "Escribe la contraseña de al menos 8 caracteres.",
            "Pulsa “Iniciar sesión”. El administrador entra al Dashboard; los demás roles entran al Inventario.",
        ],
    )
    add_callout(
        document,
        "Cuenta pendiente",
        "Si el correo todavía no fue aprobado, el sistema no permitirá entrar y mostrará que un administrador debe revisar la cuenta.",
        "warning",
    )
    add_screenshot(
        document,
        1,
        "Inicio de sesión",
        "Captura la pantalla completa con los campos Correo, Contraseña y el botón Iniciar sesión.",
    )
    add_heading(document, "3.2 Registrar una cuenta", 2)
    add_steps(
        document,
        [
            "En la pantalla de acceso, abre “Registrarse”.",
            "Escribe nombre, correo, contraseña y confirmación.",
            "Envía el formulario.",
            "Espera la aprobación del administrador. El rol inicial de las cuentas posteriores al primer usuario es Consultor.",
        ],
    )
    add_screenshot(
        document,
        2,
        "Registro de usuario",
        "Captura el formulario de registro y, si es posible, el aviso que indica que la cuenta quedó pendiente.",
    )
    add_heading(document, "3.3 Cerrar sesión", 2)
    add_body(
        document,
        "Pulsa “Cerrar sesión” en la parte inferior del menú. Hazlo siempre en computadoras compartidas. "
        "No cierres únicamente la pestaña del navegador.",
    )

    start_section(document, "4. Navegación y elementos comunes")
    add_heading(document, "4.1 Menú lateral", 2)
    add_body(
        document,
        "El menú está organizado por Inicio, Operación, Equipos, Compras, Catálogos, Herramientas, Administración "
        "y Cuenta. Los grupos pueden desplegarse y el sistema recuerda si el menú estaba expandido o compacto.",
    )
    add_screenshot(
        document,
        3,
        "Menú completo por grupos",
        "Captura el sistema como administrador con todos los grupos del menú visibles y al menos dos grupos desplegados.",
    )
    add_heading(document, "4.2 Menú compacto", 2)
    add_body(
        document,
        "Pulsa el botón de contraer para dejar solo los iconos. Coloca el cursor sobre un icono para leer su nombre "
        "en el tooltip. Vuelve a pulsar el control para ampliar el menú.",
    )
    add_screenshot(
        document,
        4,
        "Menú compacto con tooltip",
        "Captura el menú contraído y deja el cursor sobre un icono para que se vea el nombre de la opción.",
    )
    add_heading(document, "4.3 Buscador global y breadcrumbs", 2)
    add_bullets(
        document,
        [
            "Pulsa el buscador superior o usa Ctrl + K.",
            "Busca piezas, equipos, proveedores, usuarios o movimientos según tu permiso.",
            "Al abrir una pieza, el sistema la muestra dentro del inventario con el filtro y resaltado correspondientes.",
            "Los breadcrumbs indican la ruta actual, por ejemplo: Compras > Entradas pendientes > Editar.",
        ],
    )
    add_screenshot(
        document,
        5,
        "Buscador global",
        "Escribe el nombre o apodo de una pieza y captura la lista de resultados antes de abrirla.",
    )
    add_heading(document, "4.4 Favoritos, tema y notificaciones", 2)
    add_bullets(
        document,
        [
            "Usa Favoritos para fijar los apartados más utilizados.",
            "El botón de tema cambia la apariencia sin alterar datos.",
            "La campana muestra pendientes: entradas por aprobar, usuarios pendientes y alertas de stock.",
            "Los contadores rojos del menú indican acciones que requieren atención.",
        ],
    )
    add_screenshot(
        document,
        6,
        "Centro de notificaciones",
        "Captura la campana abierta mostrando contadores o actividad pendiente.",
    )
    add_heading(document, "4.5 Uso en teléfono", 2)
    add_body(
        document,
        "En celular, la navegación principal aparece en la parte inferior con accesos a Inicio, Inventario, Entrada, "
        "Salida y Menú, según el rol. Las tablas se convierten en tarjetas para evitar desplazamiento horizontal.",
    )
    add_screenshot(
        document,
        7,
        "Navegación móvil",
        "Desde un teléfono, captura una pantalla del sistema donde se vea la barra inferior y el botón Menú.",
    )

    start_section(document, "5. Dashboard gerencial")
    add_body(
        document,
        "El Dashboard es exclusivo del Administrador. Resume la situación operativa y financiera del inventario con "
        "datos de piezas reales, no con plantillas de equipos.",
    )
    add_bullets(
        document,
        [
            "Materiales registrados: número de piezas diferentes.",
            "Piezas en stock: suma de existencias actuales.",
            "Valor total del inventario: stock × costo unitario.",
            "Salidas del mes: cantidad total retirada durante el mes actual.",
            "Materiales más consumidos: clasificación de salidas del mes.",
            "Valor por categoría y estado general del stock.",
            "Alertas de stock mínimo y materiales agotados.",
            "Top 5 proveedores: monto comprado, calculado con entradas y costo unitario.",
        ],
    )
    add_callout(
        document,
        "Interpretación financiera",
        "Si un material no tiene costo unitario, aporta $0 al valor del inventario y a ciertos análisis de compras. "
        "Completa el costo desde una entrada, un XML o la edición autorizada del material.",
        "warning",
    )
    add_screenshot(
        document,
        8,
        "Dashboard gerencial completo",
        "Usa pantalla completa y captura tarjetas, gráficas, alertas y Top 5 proveedores en una sola imagen o en dos imágenes consecutivas.",
        1.65,
    )

    start_section(document, "6. Inventario de almacén")
    add_heading(document, "6.1 Consultar y localizar piezas", 2)
    add_steps(
        document,
        [
            "Abre Operación > Inventario.",
            "Busca por número de parte, código, descripción o apodo.",
            "Filtra por categoría cuando sea necesario.",
            "Confirma fotografía, categoría, almacén, descripción, apodo, marca y stock antes de operar.",
        ],
    )
    add_body(
        document,
        "El proveedor es información administrativa y puede estar oculto para el almacenista. La ubicación o almacén "
        "debe permanecer visible para facilitar el surtido físico.",
    )
    add_screenshot(
        document,
        9,
        "Inventario principal",
        "Captura la tabla con buscador, filtro de categoría, fotografías, apodo, almacén y semáforo de stock.",
    )
    add_heading(document, "6.2 Semáforo de stock", 2)
    add_table(
        document,
        ["Color", "Significado", "Acción recomendada"],
        [
            ["Verde", "Stock suficiente", "Operar normalmente."],
            ["Amarillo", "Stock cercano o igual al mínimo", "Preparar reabastecimiento."],
            ["Rojo", "Producto agotado o crítico", "No prometer salida; informar a compras."],
        ],
        [1.0, 2.5, 3.25],
    )
    add_heading(document, "6.3 Ver fotografía en grande", 2)
    add_body(
        document,
        "Pulsa la fotografía de una pieza para abrir la galería. Usa cerrar, teclado o el área exterior para regresar. "
        "La fotografía ayuda a identificar, pero no sustituye el número de parte o la medida.",
    )
    add_screenshot(
        document,
        10,
        "Visor de imagen del inventario",
        "Abre una fotografía de producto y captura el modal ampliado con el fondo atenuado.",
    )
    add_heading(document, "6.4 Controles de tabla", 2)
    add_bullets(
        document,
        [
            "Vista compacta / cómoda: cambia la densidad sin alterar datos.",
            "Columnas: muestra u oculta información para la tarea actual.",
            "Seleccionar: permite elegir varias piezas para exportar o imprimir etiquetas.",
            "Encabezado fijo: conserva los nombres de columnas al desplazarte.",
            "Paginación: cambia de página sin cargar todo el catálogo al mismo tiempo.",
        ],
    )
    add_screenshot(
        document,
        11,
        "Herramientas de tabla",
        "Captura la zona superior de la tabla con Seleccionar, Vista compacta y Columnas.",
    )

    start_section(document, "7. Registrar entradas")
    add_heading(document, "7.1 Entrada de una pieza existente", 2)
    add_steps(
        document,
        [
            "Abre Operación > Registrar entrada.",
            "Escanea el código con cámara o pistola USB, o escríbelo.",
            "Si el código existe, confirma la pieza detectada.",
            "Escribe la cantidad recibida, proveedor, costo y referencia cuando corresponda.",
            "Adjunta evidencia. Para almacenistas es obligatoria.",
            "Envía la entrada.",
        ],
    )
    add_callout(
        document,
        "Diferencia por rol",
        "La entrada del Administrador suma stock inmediatamente. La del Almacenista queda Pendiente y no modifica "
        "existencias hasta que un Administrador la apruebe.",
        "info",
    )
    add_screenshot(
        document,
        12,
        "Registrar entrada existente",
        "Escanea un código de prueba ya registrado y captura la pieza detectada, cantidad y evidencia.",
    )
    add_heading(document, "7.2 Solicitar una pieza nueva", 2)
    add_body(
        document,
        "Si el código no existe, el almacenista puede solicitar la creación de la pieza. Debe escribir al menos la "
        "descripción y la cantidad; se recomienda completar categoría, almacén, número de parte, apodo, marca, unidad, "
        "stock mínimo, stock máximo, costo, proveedor, claves SAT, fotografía del producto y evidencia de recepción.",
    )
    add_steps(
        document,
        [
            "Confirma que no exista con otra descripción, medida, categoría o apodo.",
            "Completa la descripción de la pieza y la cantidad recibida.",
            "Selecciona una categoría real del inventario y escribe la ubicación física.",
            "Adjunta una fotografía clara del producto y otra de la evidencia de recepción.",
            "Envía la solicitud para aprobación.",
        ],
    )
    add_screenshot(
        document,
        13,
        "Solicitud de material nuevo",
        "Como almacenista, captura el formulario con código nuevo, categoría desplegada, datos de la pieza y archivos de imagen.",
    )
    add_heading(document, "7.3 Aprobar, corregir o rechazar entradas", 2)
    add_body(
        document,
        "El Administrador recibe un contador en el menú y la campana. En Compras > Aprobar entradas puede filtrar "
        "pendientes, abrir evidencia, corregir datos y decidir.",
    )
    add_steps(
        document,
        [
            "Abre la entrada pendiente y verifica usuario, fecha, hora, cantidad, pieza y evidencia.",
            "Si es material nuevo, busca duplicados por número de parte, descripción, apodo, medida, categoría y código.",
            "Corrige cantidad, material vinculado, categoría, almacén, costo, proveedor o imágenes si es necesario.",
            "Pulsa Aprobar para crear o actualizar la pieza y sumar stock.",
            "Pulsa Rechazar si la entrada es incorrecta; el stock permanecerá sin cambios.",
        ],
    )
    add_screenshot(
        document,
        14,
        "Bandeja de entradas pendientes",
        "Captura la lista con usuario, fecha y hora, pieza, cantidad, estado y acciones.",
    )
    add_screenshot(
        document,
        15,
        "Corrección de entrada pendiente",
        "Abre una entrada nueva y captura categoría desplegada, datos editables, evidencia y botones Aprobar/Rechazar.",
    )
    add_screenshot(
        document,
        16,
        "Evidencia ampliada",
        "Pulsa la evidencia y captura el visor grande antes de aprobar.",
    )
    add_callout(
        document,
        "Entrada ya revisada",
        "Una solicitud aprobada o rechazada deja de ser editable. Si existe una diferencia posterior, registra el "
        "movimiento correctivo adecuado y documenta el motivo.",
        "warning",
    )

    start_section(document, "8. Registrar salidas")
    add_steps(
        document,
        [
            "Abre Operación > Registrar salida.",
            "Escanea el código con cámara, pistola USB o selección manual.",
            "Si el mismo código aparece en varias piezas, busca manualmente y elige la categoría y almacén correctos.",
            "Escribe la cantidad a retirar.",
            "Agrega referencia u orden y notas de quién lo pidió o para qué se utilizará.",
            "Revisa el producto seleccionado y pulsa Registrar salida.",
        ],
    )
    add_callout(
        document,
        "Protección de stock",
        "El sistema no permite retirar más piezas de las disponibles. Una salida aprobada descuenta stock de inmediato "
        "y queda asociada al usuario, fecha, hora, referencia y motivo.",
        "danger",
    )
    add_screenshot(
        document,
        17,
        "Registrar salida",
        "Captura el formulario con código/cámara, cantidad, referencia, notas y una tarjeta de pieza seleccionada con categoría y almacén.",
    )
    add_heading(document, "8.1 Lectores compatibles", 2)
    add_bullets(
        document,
        [
            "Pistola USB: funciona como teclado. Coloca el cursor en Código de barras, apunta y dispara.",
            "Cámara: concede permiso y enfoca el código dentro del marco.",
            "Imagen guardada: selecciona una foto clara del código cuando la cámara no esté disponible.",
            "QR interno: se lee igual que un código y abre la pieza vinculada.",
        ],
    )

    start_section(document, "9. Devoluciones y mermas")
    add_heading(document, "9.1 Devolución", 2)
    add_body(
        document,
        "Se usa cuando una pieza salió del almacén, no fue utilizada y regresa en condiciones aptas. La devolución "
        "aumenta el stock y conserva el motivo y la referencia.",
    )
    add_heading(document, "9.2 Merma o scrap", 2)
    add_body(
        document,
        "Se usa para retirar definitivamente piezas dañadas, defectuosas o inutilizables. La merma reduce el stock y "
        "exige una fotografía de evidencia. No puede superar la existencia disponible.",
    )
    add_steps(
        document,
        [
            "Abre Operación > Devoluciones y mermas.",
            "Busca y selecciona la pieza exacta.",
            "Elige Devolución o Merma.",
            "Escribe cantidad, referencia y motivo.",
            "Para Merma, adjunta una fotografía clara del daño.",
            "Confirma el movimiento y revisa el stock resultante.",
        ],
    )
    add_screenshot(
        document,
        18,
        "Devolución o merma",
        "Captura el formulario con tipo de movimiento, pieza, cantidad, motivo y campo de evidencia.",
    )

    start_section(document, "10. Equipos y paquetes")
    add_heading(document, "10.1 Concepto", 2)
    add_body(
        document,
        "Un equipo o paquete es una receta de materiales. Cada renglón indica una pieza real del inventario y la "
        "cantidad que consume una unidad del equipo. Retirar o vender un equipo descuenta automáticamente todas las piezas vinculadas.",
    )
    add_heading(document, "10.2 Crear un equipo", 2)
    add_steps(
        document,
        [
            "Abre Equipos > Equipos y paquetes.",
            "Escribe nombre, código opcional y descripción.",
            "Guarda el equipo y ábrelo para configurar sus piezas.",
        ],
    )
    add_screenshot(
        document,
        19,
        "Listado y alta de equipos",
        "Captura la pantalla con el formulario de nuevo equipo y la lista de equipos existentes.",
    )
    add_heading(document, "10.3 Agregar piezas al equipo", 2)
    add_steps(
        document,
        [
            "En el detalle del equipo, selecciona una pieza real del inventario.",
            "Confirma la fotografía y los datos que se llenan automáticamente.",
            "Escribe solamente la cantidad por paquete y, si hace falta, una nota.",
            "Agrega la pieza y repite hasta completar la receta.",
            "No dejes renglones “Sin vincular” si el equipo se utilizará para descontar stock.",
        ],
    )
    add_screenshot(
        document,
        20,
        "Detalle de equipo con fotografías",
        "Captura Piezas requeridas y Agregar pieza; debe verse la foto de la pieza seleccionada y la cantidad por paquete.",
    )
    add_heading(document, "10.4 Retirar o vender equipos", 2)
    add_steps(
        document,
        [
            "Abre Equipos > Retirar equipo.",
            "Busca y selecciona el equipo.",
            "Elige Venta o Retiro.",
            "Indica cuántos paquetes se retiran, referencia y notas.",
            "Revisa el resumen de piezas necesarias y confirma.",
        ],
    )
    add_callout(
        document,
        "Validación automática",
        "El sistema bloquea el retiro si existe una pieza sin vincular o si falta stock. El aviso identifica la pieza, "
        "la cantidad disponible, la requerida y la faltante. Corrige la receta o repón inventario antes de continuar.",
        "danger",
    )
    add_screenshot(
        document,
        21,
        "Retiro o venta de equipo",
        "Selecciona un equipo de prueba y captura tipo, cantidad, resumen de componentes y botón de confirmación.",
    )
    add_screenshot(
        document,
        22,
        "Aviso de pieza faltante",
        "En pruebas, intenta retirar una cantidad que exceda un componente y captura el mensaje con la pieza faltante.",
    )
    add_heading(document, "10.5 Historial de equipos", 2)
    add_body(
        document,
        "El historial muestra ventas y retiros con fecha, usuario, cantidad, referencia, notas y piezas descontadas. "
        "Úsalo para rastrear por qué bajó el inventario.",
    )
    add_screenshot(
        document,
        23,
        "Historial de equipos",
        "Captura la lista filtrada por Venta o Retiro, mostrando usuario, fecha y componentes.",
    )

    start_section(document, "11. Compras y facturas")
    add_heading(document, "11.1 Importar factura XML", 2)
    add_body(
        document,
        "La importación acepta CFDI XML con UUID del SAT. Lee proveedor, RFC, serie, folio, fecha, moneda, subtotal, "
        "descuento, impuestos, total, forma y método de pago, además de los conceptos con claves SAT, unidades, "
        "número de parte, cantidad, precio unitario e importes.",
    )
    add_steps(
        document,
        [
            "Abre Compras > Importar XML.",
            "Selecciona el archivo .xml y pulsa Vista previa.",
            "Revisa encabezado de factura, proveedor, UUID, totales y conceptos.",
            "Marca únicamente los conceptos que sí representan piezas enteras de inventario.",
            "Asigna una categoría real a cada concepto nuevo.",
            "Confirma la importación.",
        ],
    )
    add_screenshot(
        document,
        24,
        "Carga de factura XML",
        "Captura la pantalla para seleccionar el archivo XML y el botón de vista previa.",
    )
    add_screenshot(
        document,
        25,
        "Vista previa de factura XML",
        "Captura los datos fiscales superiores y la tabla de conceptos con selección y categoría.",
        1.6,
    )
    add_callout(
        document,
        "Cómo afecta el inventario",
        "Si el número de parte ya existe, el XML suma la cantidad y actualiza costo, proveedor y datos fiscales. "
        "Si no existe, crea una pieza nueva sin código de barras. Las cantidades decimales no se importan porque el inventario usa piezas enteras.",
        "purple",
    )
    add_callout(
        document,
        "Factura repetida",
        "Un XML ya registrado puede volver a previsualizarse para consulta. El sistema muestra un aviso y bloquea una "
        "segunda importación para evitar duplicar stock.",
        "warning",
    )
    add_screenshot(
        document,
        26,
        "Aviso de XML ya importado",
        "Vuelve a previsualizar un XML de prueba ya registrado y captura el aviso de duplicado sin importarlo nuevamente.",
    )
    add_heading(document, "11.2 Proveedores", 2)
    add_body(
        document,
        "Compras > Proveedores agrupa piezas por nombre de proveedor y muestra RFC disponible, cantidad de productos, "
        "piezas y valor. Al abrir un proveedor se consultan todos los materiales asociados.",
    )
    add_screenshot(
        document,
        27,
        "Listado de proveedores",
        "Captura proveedores con RFC, productos, piezas y valor.",
    )
    add_screenshot(
        document,
        28,
        "Detalle de proveedor",
        "Abre un proveedor y captura la lista de piezas vendidas o asociadas a ese proveedor.",
    )
    add_heading(document, "11.3 Órdenes de compra", 2)
    add_body(
        document,
        "La orden de compra planea una adquisición. Incluye proveedor, referencia, fecha de orden, fecha esperada, "
        "notas y uno o más renglones con material, descripción, cantidad, costo y subtotal.",
    )
    add_steps(
        document,
        [
            "Abre Compras > Órdenes de compra.",
            "Selecciona proveedor y fechas.",
            "Agrega materiales o descripciones libres con cantidad y costo.",
            "Guarda como Borrador.",
            "Actualiza el estado a Enviada, Recibida o Cancelada según avance.",
        ],
    )
    add_callout(
        document,
        "Importante",
        "Cambiar una orden a Recibida no suma existencias. Registra la entrada real con evidencia y aprobación para modificar stock.",
        "danger",
    )
    add_screenshot(
        document,
        29,
        "Órdenes de compra",
        "Captura el formulario con dos renglones de ejemplo y la lista de órdenes con sus estados.",
    )

    start_section(document, "12. Catálogos y etiquetas")
    add_heading(document, "12.1 Categorías", 2)
    add_body(
        document,
        "Las categorías describen tipos de piezas reales. Los nombres de equipos no son categorías. El Administrador "
        "puede crear, renombrar, describir, activar o desactivar una categoría.",
    )
    add_bullets(
        document,
        [
            "Al renombrar una categoría, los materiales asociados se actualizan.",
            "Una categoría en uso no se elimina; puede desactivarse.",
            "No crees duplicados por mayúsculas, acentos o abreviaturas.",
        ],
    )
    add_screenshot(
        document,
        30,
        "Administración de categorías",
        "Captura Nueva categoría y la tabla de categorías con materiales, edición y estado.",
    )
    add_heading(document, "12.2 Materiales / catálogo completo", 2)
    add_body(
        document,
        "El catálogo completo es exclusivo del Administrador y muestra todos los datos de cada pieza: fotografía, "
        "categoría, almacén, número de parte, código, apodo, descripción, marca, proveedor, claves SAT, stock, mínimos, máximos, costo y moneda.",
    )
    add_screenshot(
        document,
        31,
        "Catálogo completo",
        "Captura una página con fotografías y columnas administrativas; usa una pieza de prueba sin información sensible.",
    )
    add_heading(document, "12.3 Editar o eliminar materiales", 2)
    add_bullets(
        document,
        [
            "Editar permite corregir datos, foto, código, categoría, ubicación, costos y existencias.",
            "Para movimientos normales de stock usa Entrada, Salida, Devolución o Merma; evita cambiar stock desde Editar.",
            "Eliminar borra la pieza y puede afectar trazabilidad. Úsalo solo para registros creados por error y con autorización.",
        ],
    )
    add_callout(
        document,
        "Correcciones de inventario",
        "Un cambio directo de stock desde Editar queda como edición de catálogo, no como movimiento de entrada o salida. "
        "Para auditoría clara, utiliza el módulo operativo correspondiente.",
        "warning",
    )
    add_heading(document, "12.4 Historial de una pieza", 2)
    add_body(
        document,
        "Desde el catálogo completo, abre Historial para consultar altas, ediciones y movimientos de la pieza, con usuario y fecha.",
    )
    add_screenshot(
        document,
        32,
        "Historial de material",
        "Abre una pieza con varios movimientos y captura la línea de tiempo o tabla de acciones.",
    )
    add_heading(document, "12.5 Códigos y etiquetas", 2)
    add_bullets(
        document,
        [
            "Agregar código: escanea o escribe el código de barras del fabricante.",
            "Generar QR: crea un identificador interno único cuando la pieza no tiene código.",
            "Impresión por lote: selecciona varias piezas y genera etiquetas juntas.",
            "Un código no puede pertenecer a dos materiales distintos.",
        ],
    )
    add_screenshot(
        document,
        33,
        "Agregar códigos y generar etiquetas",
        "Captura la lista de piezas sin código, el control para escanear y la acción para generar QR.",
    )
    add_screenshot(
        document,
        34,
        "Etiqueta lista para imprimir",
        "Abre una etiqueta y captura logo, nombre de pieza, número de parte y QR/código.",
    )

    start_section(document, "13. Herramientas y reportes")
    add_heading(document, "13.1 Identificador visual", 2)
    add_body(
        document,
        "El identificador compara una foto tomada o subida con fotografías reales del inventario. No requiere escribir "
        "medidas. Los resultados muestran coincidencia aproximada y datos para confirmar.",
    )
    add_steps(
        document,
        [
            "Abre Herramientas > Identificador visual.",
            "Pulsa Tomar foto o Subir imagen.",
            "Centra una sola pieza, con buena luz y fondo limpio. Evita manos, herramientas y objetos alrededor.",
            "Espera las sugerencias y compara fotografía, categoría, descripción, número de parte y medida.",
            "Pulsa Ver para abrir la coincidencia en el inventario.",
        ],
    )
    add_callout(
        document,
        "La puntuación no es una garantía",
        "Una foto de celular con fondo ocupado, perspectiva fuerte o mala iluminación puede producir sugerencias fuera "
        "de lugar. Nunca registres un movimiento solamente por la puntuación visual.",
        "warning",
    )
    add_screenshot(
        document,
        35,
        "Tomar o subir foto",
        "Desde un teléfono, captura la pantalla del identificador con los botones de cámara y archivo.",
    )
    add_screenshot(
        document,
        36,
        "Resultados del identificador visual",
        "Analiza una pieza de prueba y captura la foto procesada y las sugerencias con puntaje.",
    )
    add_heading(document, "13.2 Reportes", 2)
    add_body(
        document,
        "El centro de reportes está disponible para Administrador y Consultor. Resume materiales, stock total, valor "
        "del inventario y movimientos recientes.",
    )
    add_bullets(
        document,
        [
            "Inventario CSV/Excel: exporta datos para análisis.",
            "Inventario PDF: genera un reporte visual con logo.",
            "Salidas CSV/Excel: exporta fecha, pieza, cantidad, stock anterior/nuevo, referencia, motivo y usuario.",
            "Salidas PDF: genera un documento de auditoría con hasta los movimientos más recientes definidos por el sistema.",
            "En Inventario, la selección múltiple permite exportar solo las piezas elegidas.",
        ],
    )
    add_screenshot(
        document,
        37,
        "Centro de reportes",
        "Captura indicadores, movimientos recientes y botones de exportación.",
    )

    start_section(document, "14. Administración")
    add_heading(document, "14.1 Usuarios y permisos", 2)
    add_steps(
        document,
        [
            "Abre Administración > Usuarios.",
            "Localiza la cuenta pendiente.",
            "Selecciona Administrador, Almacenista o Consultor.",
            "Marca Aprobar correo.",
            "Pulsa Guardar permisos.",
        ],
    )
    add_body(
        document,
        "Para retirar acceso, desmarca la aprobación o cambia el rol según la política interna. No compartas cuentas: "
        "cada persona debe usar su propio correo para conservar la auditoría.",
    )
    add_screenshot(
        document,
        38,
        "Usuarios y permisos",
        "Captura una cuenta pendiente y otra aprobada, mostrando rol, estado y Guardar permisos.",
    )
    add_heading(document, "14.2 Auditoría", 2)
    add_body(
        document,
        "La auditoría traduce a lenguaje sencillo quién realizó una acción, en qué módulo, cuándo y con qué detalle. "
        "Sirve para investigar altas, cambios, movimientos, XML, equipos, usuarios y respaldos.",
    )
    add_callout(
        document,
        "Eliminar registros de auditoría",
        "Borrar un evento o limpiar el historial es irreversible. Hazlo únicamente con autorización y después de generar el respaldo correspondiente.",
        "danger",
    )
    add_screenshot(
        document,
        39,
        "Registro de auditoría",
        "Captura filtros, usuario, fecha, módulo, acción y descripción entendible.",
    )
    add_heading(document, "14.3 Historial de salidas", 2)
    add_body(
        document,
        "Muestra las salidas de almacén realizadas, con fecha, producto, cantidad, usuario, referencia, motivo y stock antes/después.",
    )
    add_screenshot(
        document,
        40,
        "Historial administrativo de salidas",
        "Captura varias salidas mostrando usuario, fecha, cantidad y referencia.",
    )
    add_heading(document, "14.4 Respaldos", 2)
    add_steps(
        document,
        [
            "Abre Administración > Respaldos.",
            "Pulsa Generar respaldo y conserva el archivo .sql descargado en una ubicación protegida.",
            "Antes de restaurar, genera un respaldo del estado actual.",
            "Selecciona el archivo .sql autorizado y confirma la restauración durante una ventana de mantenimiento.",
        ],
    )
    add_callout(
        document,
        "Alcance del respaldo",
        "El archivo SQL incluye la base de datos, pero no incluye las fotografías almacenadas como archivos locales. "
        "La empresa debe respaldar también la carpeta de almacenamiento público de imágenes.",
        "danger",
    )
    add_callout(
        document,
        "Restaurar reemplaza datos",
        "Una restauración puede sobrescribir el estado actual. No la ejecutes con usuarios operando ni sin verificar el archivo.",
        "warning",
    )
    add_screenshot(
        document,
        41,
        "Respaldos de base de datos",
        "Captura lista de respaldos, botón para generar y sección para restaurar archivo SQL.",
    )

    start_section(document, "15. Perfil, seguridad y dispositivos")
    add_heading(document, "15.1 Perfil", 2)
    add_bullets(
        document,
        [
            "Actualizar nombre y correo.",
            "Cambiar contraseña usando la contraseña actual y una nueva de al menos 8 caracteres.",
            "Consultar actividad reciente propia.",
            "Eliminar cuenta solo con autorización; nunca elimines al último administrador.",
        ],
    )
    add_screenshot(
        document,
        42,
        "Perfil del usuario",
        "Captura las secciones de información, cambio de contraseña y actividad reciente.",
    )
    add_heading(document, "15.2 Cámara en celular", 2)
    add_bullets(
        document,
        [
            "Acepta el permiso de cámara cuando el navegador lo solicite.",
            "En iPhone, revisa Ajustes > Chrome o Safari > Cámara y permite el acceso.",
            "Si la cámara se bloquea por seguridad de la conexión, utiliza Subir imagen y toma la foto desde el selector del teléfono.",
            "Limpia la lente, usa buena luz, centra el objeto y evita el zoom digital excesivo.",
        ],
    )
    add_heading(document, "15.3 Pistola USB", 2)
    add_steps(
        document,
        [
            "Conecta la pistola al puerto USB.",
            "Abre una pantalla con campo Código de barras.",
            "Haz clic dentro del campo.",
            "Escanea. La pistola escribe el código como teclado y normalmente envía Enter.",
            "Confirma siempre la pieza detectada antes de guardar.",
        ],
    )

    start_section(document, "16. Alertas y notificaciones")
    add_heading(document, "16.1 Alertas dentro del sistema", 2)
    add_body(
        document,
        "El sistema muestra alertas por stock mínimo, piezas agotadas, entradas pendientes y usuarios por aprobar. "
        "El Administrador debe revisar la campana y los contadores al iniciar su jornada.",
    )
    add_heading(document, "16.2 Correo automático de stock", 2)
    add_body(
        document,
        "Cuando el servicio programado está configurado, todos los días a las 08:00, hora de Ciudad de México, el sistema "
        "revisa materiales con stock menor o igual al mínimo y envía un correo al destinatario definido por la empresa.",
    )
    add_callout(
        document,
        "Requisito operativo",
        "El correo automático depende de la configuración del servidor, el destinatario de alertas y el programador de "
        "tareas. Si no llegan mensajes, las alertas visuales del sistema siguen disponibles.",
        "info",
    )

    start_section(document, "17. Reglas de stock y trazabilidad")
    add_table(
        document,
        ["Operación", "Efecto en stock", "Cuándo se aplica"],
        [
            ["Entrada por administrador", "Suma", "Al guardar."],
            ["Entrada por almacenista", "No cambia al enviar", "Suma cuando el administrador aprueba."],
            ["Importación XML", "Suma", "Al confirmar una factura no duplicada."],
            ["Salida de pieza", "Resta", "Al registrar la salida."],
            ["Devolución", "Suma", "Al confirmar la devolución."],
            ["Merma", "Resta", "Al confirmar con evidencia."],
            ["Venta/retiro de equipo", "Resta componentes", "Al validar todos los vínculos y existencias."],
            ["Orden de compra", "No cambia", "La recepción debe registrarse como entrada."],
            ["Editar material", "Puede corregir", "Solo administrador; evitarlo para operación normal."],
        ],
        [2.0, 1.65, 3.1],
    )
    add_heading(document, "17.1 Datos mínimos recomendados", 2)
    add_bullets(
        document,
        [
            "Descripción clara que incluya pieza, medida y característica importante.",
            "Número de parte único cuando exista.",
            "Apodo usado por el personal.",
            "Categoría real y almacén o ubicación física.",
            "Fotografía frontal y limpia.",
            "Marca, unidad, costo y proveedor cuando se conozcan.",
            "Stock mínimo y máximo.",
            "Código de barras del fabricante o QR interno.",
        ],
    )

    start_section(document, "18. Rutinas recomendadas")
    add_heading(document, "18.1 Inicio de jornada del administrador", 2)
    add_steps(
        document,
        [
            "Revisar campana, usuarios y entradas pendientes.",
            "Aprobar o rechazar entradas después de verificar evidencia.",
            "Revisar alertas de stock, valor y Top 5 proveedores.",
            "Atender piezas agotadas o por debajo del mínimo.",
            "Revisar auditoría cuando exista una diferencia.",
        ],
    )
    add_heading(document, "18.2 Inicio de jornada del almacenista", 2)
    add_steps(
        document,
        [
            "Consultar inventario y ubicaciones.",
            "Registrar toda recepción con evidencia.",
            "No entregar piezas sin registrar salida.",
            "Registrar devoluciones y mermas el mismo día.",
            "Confirmar piezas y cantidades antes de retirar equipos.",
        ],
    )
    add_heading(document, "18.3 Revisión semanal", 2)
    add_bullets(
        document,
        [
            "Verificar piezas sin fotografía, código, costo, ubicación o stock mínimo.",
            "Comparar existencias críticas con compras pendientes.",
            "Generar reportes de inventario y salidas.",
            "Crear respaldo SQL y copia de fotografías.",
            "Revisar cuentas activas y permisos.",
        ],
    )

    start_section(document, "19. Solución de problemas")
    add_table(
        document,
        ["Situación", "Qué revisar"],
        [
            ["No puedo iniciar sesión", "Correo correcto, contraseña de 8 caracteres y cuenta aprobada."],
            ["No veo una opción del menú", "El rol no tiene permiso o el grupo está plegado."],
            ["La cámara no abre", "Permiso del navegador; usa Subir imagen como alternativa."],
            ["La pistola no escanea", "Cursor dentro del campo, distribución de teclado y Enter configurado."],
            ["No se ven fotografías", "Verificar almacenamiento de imágenes y vínculo público en esa instalación."],
            ["Una entrada no sumó stock", "Si la creó un almacenista, debe aprobarla el administrador."],
            ["No puedo importar XML", "Archivo CFDI válido, UUID, conceptos con NoIdentificacion y cantidades enteras."],
            ["XML ya registrado", "Puede previsualizarse, pero no importarse de nuevo."],
            ["No puedo retirar equipo", "Hay piezas sin vincular o stock insuficiente; leer el detalle del aviso."],
            ["El valor aparece en $0", "Completar costo unitario y moneda de los materiales."],
            ["Una foto da coincidencias extrañas", "Fondo limpio, una sola pieza, buena luz y confirmar datos manualmente."],
            ["La hora no coincide", "Reportar al administrador para revisar zona horaria del sistema/servidor."],
        ],
        [2.25, 4.5],
    )

    start_section(document, "20. Glosario")
    add_table(
        document,
        ["Término", "Definición"],
        [
            ["Pieza o material", "Existencia física individual controlada por el inventario."],
            ["Equipo o paquete", "Conjunto o receta que consume cantidades definidas de piezas reales."],
            ["Stock", "Cantidad disponible actualmente."],
            ["Stock mínimo", "Nivel que activa una alerta de reabastecimiento."],
            ["Stock máximo", "Límite recomendado de almacenamiento."],
            ["Entrada pendiente", "Recepción capturada por un almacenista que espera aprobación."],
            ["Merma / scrap", "Baja definitiva por daño, defecto o inutilidad."],
            ["Devolución", "Regreso de una pieza utilizable al inventario."],
            ["UUID", "Identificador fiscal único de una factura CFDI."],
            ["Clave SAT", "Clave de producto o servicio incluida en el CFDI."],
            ["QR interno", "Código generado por el sistema para piezas sin código del fabricante."],
            ["Auditoría", "Registro de acciones realizadas por los usuarios."],
        ],
        [1.7, 5.05],
    )

    start_section(document, "Anexo A. Lista maestra de capturas")
    add_body(
        document,
        "Toma las capturas en el orden siguiente. Usa una cuenta Administrador para las pantallas administrativas, "
        "una cuenta Almacenista para demostrar aprobaciones y un teléfono para las pantallas móviles.",
    )
    screenshot_rows = [
        ["01", "Inicio de sesión", "Pantalla completa, sin contraseña visible."],
        ["02", "Registro", "Formulario y aviso de cuenta pendiente."],
        ["03", "Menú completo", "Administrador; grupos visibles."],
        ["04", "Menú compacto", "Iconos y tooltip."],
        ["05", "Buscador global", "Resultados de una pieza por apodo."],
        ["06", "Notificaciones", "Campana abierta."],
        ["07", "Navegación móvil", "Barra inferior y Menú."],
        ["08", "Dashboard", "Tarjetas, gráficas, alertas y proveedores."],
        ["09", "Inventario", "Tabla con foto, apodo, almacén y stock."],
        ["10", "Foto ampliada", "Modal de imagen."],
        ["11", "Controles de tabla", "Densidad, columnas y selección."],
        ["12", "Entrada existente", "Pieza detectada y evidencia."],
        ["13", "Material nuevo", "Formulario de almacenista."],
        ["14", "Pendientes", "Bandeja del administrador."],
        ["15", "Editar pendiente", "Datos, categoría y decisiones."],
        ["16", "Evidencia", "Imagen abierta en grande."],
        ["17", "Salida", "Selección exacta y cantidad."],
        ["18", "Devolución/merma", "Tipo, motivo y evidencia."],
        ["19", "Equipos", "Alta y listado."],
        ["20", "Detalle de equipo", "Piezas con fotografías."],
        ["21", "Retirar equipo", "Resumen de componentes."],
        ["22", "Stock insuficiente", "Aviso de pieza faltante."],
        ["23", "Historial de equipos", "Venta o retiro."],
        ["24", "Subir XML", "Selector de archivo."],
        ["25", "Vista previa XML", "Datos fiscales y conceptos."],
        ["26", "XML duplicado", "Aviso sin volver a importar."],
        ["27", "Proveedores", "Resumen por proveedor."],
        ["28", "Detalle proveedor", "Piezas asociadas."],
        ["29", "Órdenes de compra", "Formulario y estados."],
        ["30", "Categorías", "Alta y tabla."],
        ["31", "Catálogo completo", "Foto y datos administrativos."],
        ["32", "Historial de material", "Movimientos y cambios."],
        ["33", "Códigos", "Agregar código y generar QR."],
        ["34", "Etiqueta", "Vista lista para imprimir."],
        ["35", "Identificador móvil", "Tomar o subir foto."],
        ["36", "Resultados visuales", "Sugerencias y puntajes."],
        ["37", "Reportes", "Indicadores y descargas."],
        ["38", "Usuarios", "Aprobación y roles."],
        ["39", "Auditoría", "Acciones entendibles."],
        ["40", "Historial de salidas", "Usuario, fecha y cantidad."],
        ["41", "Respaldos", "Generar y restaurar."],
        ["42", "Perfil", "Datos, contraseña y actividad."],
    ]
    add_table(document, ["No.", "Pantalla", "Qué debe verse"], screenshot_rows, [0.55, 2.15, 4.05])
    add_callout(
        document,
        "Privacidad antes de entregar",
        "Difumina correos personales, RFC, UUID, montos sensibles, referencias internas y cualquier fotografía de "
        "personas. No muestres contraseñas, tokens ni archivos reales de clientes.",
        "danger",
    )

    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
