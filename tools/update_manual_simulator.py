from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / 'docs' / 'Manual_de_Usuario_Inventario_Lugarth.docx'
OUTPUT = ROOT / 'docs' / 'Manual_de_Usuario_Inventario_Lugarth_Actualizado.docx'


def insert_paragraph_after(paragraph, text, style):
    element = OxmlElement('w:p')
    paragraph._p.addnext(element)
    created = Paragraph(element, paragraph._parent)
    created.style = style
    created.add_run(text)
    return created


def update_capture_table(table, title, instruction):
    paragraphs = table.cell(0, 0).paragraphs
    paragraphs[0].runs[0].text = title
    paragraphs[1].runs[0].text = instruction
    paragraphs[2].runs[0].text = 'Inserta la imagen en este espacio y elimina este recuadro.'


def main():
    document = Document(SOURCE)

    # Keep the document control information current.
    document.tables[1].cell(0, 1).text = '1.1'
    document.tables[1].cell(1, 1).text = date(2026, 8, 6).strftime('%d/%m/%Y')

    # Add the new annex to the existing contents grid.
    contents = document.tables[4]
    contents_row = contents.add_row().cells
    contents_row[0].text = ''
    contents_row[1].text = ''
    contents_row[2].text = 'Anexo B'
    contents_row[3].text = 'Capturas del simulador de equipos'

    history_heading = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip() == '10.5 Historial de equipos'
    )
    history_text = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith('El historial muestra ventas y retiros')
    )

    anchor = history_text
    anchor = insert_paragraph_after(anchor, '10.6 Simulador de fabricacion', 'Heading 2')
    anchor = insert_paragraph_after(
        anchor,
        'El simulador indica cuantos equipos completos se pueden fabricar o vender con el stock real actual. No reserva ni descuenta piezas: sirve para planear antes de registrar una venta o un retiro.',
        'Normal',
    )
    for step in [
        'Abre Equipos > Simulador de equipos.',
        'Selecciona el equipo que quieres revisar.',
        'Confirma la tarjeta Equipos fabricables hoy. Ese numero depende de la pieza con menor capacidad.',
        'Revisa Pieza limitante para saber cual material detiene la fabricacion primero.',
        'Escribe la cantidad que deseas fabricar o vender. La tabla muestra por cada pieza el stock, lo requerido, lo que quedaria y, si aplica, lo que falta.',
        'Si el resultado indica stock completo, usa Retirar o vender para realizar el movimiento real. El simulador por si solo no modifica existencias.',
    ]:
        anchor = insert_paragraph_after(anchor, step, 'List Number')
    insert_paragraph_after(
        anchor,
        'Importante: la capacidad se calcula para un solo tipo de equipo a la vez. Si planeas fabricar equipos diferentes que comparten piezas, revisa cada equipo antes de confirmar movimientos.',
        'Normal',
    )

    # Append screenshot slots using the same styled capture box used by the rest of the manual.
    document.add_page_break()
    document.add_heading('Anexo B. Capturas del simulador de equipos', level=1)
    document.add_paragraph('Agrega estas tres capturas despues de validar el simulador con datos reales. Muestra solo informacion de prueba o difumina datos sensibles.', style='Normal')
    template = document.tables[36]._tbl
    captures = [
        (
            'CAPTURA 43  |  Acceso al simulador',
            'Abre Equipos > Simulador de equipos y captura el selector de equipo junto con la lista de equipos disponibles.',
        ),
        (
            'CAPTURA 44  |  Capacidad y pieza limitante',
            'Selecciona un equipo con receta completa y captura Equipos fabricables hoy, costo aproximado y la tarjeta de pieza limitante.',
        ),
        (
            'CAPTURA 45  |  Simulacion de faltantes',
            'Escribe una cantidad mayor a la capacidad disponible y captura la tabla resaltando la pieza faltante, cantidad requerida y faltante exacto.',
        ),
    ]
    for title, instruction in captures:
        table_element = deepcopy(template)
        body = document._body._element
        body.insert(body.index(body.sectPr), table_element)
        capture_table = Table(table_element, document._body)
        update_capture_table(capture_table, title, instruction)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()
